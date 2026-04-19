import sys
try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

doc = Document(r'requerimientos\Informe_Servidores.docx')

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    print(f"[{i}] Style={p.style.name}: {p.text}")

print("\n=== TABLES ===")
for i, t in enumerate(doc.tables):
    print(f"\n--- Table {i} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
    for j, row in enumerate(t.rows):
        cells = [cell.text.replace('\n', ' | ') for cell in row.cells]
        print(f"  Row {j}: {' || '.join(cells)}")

print("\n=== SECTIONS ===")
for i, section in enumerate(doc.sections):
    print(f"Section {i}: width={section.page_width}, height={section.page_height}")
    print(f"  Margins: top={section.top_margin}, bottom={section.bottom_margin}, left={section.left_margin}, right={section.right_margin}")

print("\n=== IMAGES ===")
from docx.opc.constants import RELATIONSHIP_TYPE as RT
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        print(f"Image: {rel.target_ref}")
